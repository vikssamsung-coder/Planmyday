"""
report.py — builds the daily report workbook the user downloads and emails.

This file is the CONTRACT between the User app and the Master app. The Master
app ingests these workbooks (from an email folder) to build its dashboard, so
the sheet names and column order here are an API — change them only on purpose,
and bump REPORT_VERSION when you do.

Two readers, two needs, one file:
  - a human opens it before emailing  -> the "Summary" sheet is formatted & readable
  - the Master app parses it           -> the data sheets carry stable columns + VALUES

Why values, not formulas: openpyxl writes formulas with no cached result, and the
running app has no Excel/LibreOffice to recalc. The Master app reads with pandas,
which would get blank cells. A report is a frozen snapshot anyway, so we freeze
computed numbers as values.

Filename convention (the Master app matches on this):
    PlanMyDay_<user_key>_<YYYY-MM-DD>.xlsx
"""

import datetime as dt
import io
import os

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

import storage
import nudge
import paths


def _fmt(v):
    """Trim trailing .0 from whole numbers, add thousands separators. Leaves text alone."""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return v
    if f == int(f):
        return f"{int(f):,}"
    return f"{f:,.1f}"

REPORT_VERSION = "1.0"
APP_VERSION = "0.2"

# ---- styling -------------------------------------------------------------
FONT = "Arial"
INK = "1A1A1A"
MUTED = "6B6B6B"
HEAD_FILL = "2D4A5E"          # deep slate header
BAND = "EEF2F5"               # zebra band
GOOD = "0F6E56"
WARN = "BA7517"
BAD = "A32D2D"

_thin = Side(style="thin", color="D9D9D9")
BORDER = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)


def _status_color(status):
    return {
        "Ahead": GOOD, "On Track": GOOD,
        "Behind": WARN, "Critical": BAD, "No Target": MUTED,
    }.get(status, INK)


# ---- data assembly -------------------------------------------------------

def _scorecard_rows(user_key, month, on_date):
    """One scorecard dict per KPI, snapshot as of on_date."""
    targets = storage.get_targets(user_key, month)
    rows = []
    for _, r in targets.iterrows():
        s = nudge.score_kpi(
            r["monthly_target"], r["achieved_mtd"],
            on_date.year, on_date.month, today=on_date,
        )
        s["kpi_name"] = r["kpi_name"]
        s["target_unit"] = r.get("target_unit", "")
        s["priority"] = r.get("priority", "")
        rows.append(s)
    return rows


def _task_rows(user_key, date_str):
    tasks = storage.get_tasks(user_key, date_str)
    return tasks.to_dict("records") if not tasks.empty else []


def _day_update(user_key, date_str):
    ups = storage.get_day_updates(user_key, date_str)
    if ups.empty:
        return {}
    return ups.iloc[-1].to_dict()


# ---- sheet writers -------------------------------------------------------

def _write_table(ws, headers, rows, start_row=1):
    """Generic banded table with a slate header. Returns next free row."""
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=start_row, column=c, value=h)
        cell.font = Font(name=FONT, bold=True, color="FFFFFF", size=10)
        cell.fill = PatternFill("solid", fgColor=HEAD_FILL)
        cell.alignment = Alignment(horizontal="left", vertical="center")
        cell.border = BORDER
    for i, row in enumerate(rows, start=1):
        band = PatternFill("solid", fgColor=BAND) if i % 2 == 0 else None
        for c, h in enumerate(headers, start=1):
            v = row.get(h, "")
            cell = ws.cell(row=start_row + i, column=c, value=v)
            cell.font = Font(name=FONT, size=10, color=INK)
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=False)
            cell.border = BORDER
            if band:
                cell.fill = band
    return start_row + len(rows) + 1


def _autofit(ws, max_w=48):
    for col in ws.columns:
        width = 10
        letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value is not None:
                width = max(width, min(len(str(cell.value)) + 2, max_w))
        ws.column_dimensions[letter].width = width


# ---- the contract sheets -------------------------------------------------

META_COLS = ["report_version", "app_version", "user_key", "name", "role",
             "department", "report_date", "generated_at", "month"]

SCORECARD_COLS = ["kpi_name", "monthly_target", "achieved_mtd", "target_unit",
                  "achievement_pct", "expected_pct", "gap",
                  "required_run_rate", "current_run_rate",
                  "remaining_working_days", "status", "priority"]

TASK_COLS = ["task_id", "title", "category", "priority", "horizon",
             "goal_aligned", "alignment_note", "linked_kpi",
             "expected_output", "success_metric", "status", "done_at", "notes"]

UPDATE_COLS = ["completed_tasks", "pending_tasks", "blocked_tasks",
               "numbers_update", "what_worked", "what_did_not_work", "remarks"]


def build_report(user_key, report_date=None):
    """Build the report workbook. Returns (filename, bytes)."""
    user = storage.get_user(user_key)
    if user is None:
        raise ValueError(f"unknown user: {user_key}")

    on_date = report_date or dt.date.today()
    if isinstance(on_date, str):
        on_date = dt.date.fromisoformat(on_date)
    date_str = on_date.isoformat()
    month = on_date.strftime("%Y-%m")

    scards = _scorecard_rows(user_key, month, on_date)
    tasks = _task_rows(user_key, date_str)
    upd = _day_update(user_key, date_str)

    wb = Workbook()

    # ---------- Summary (human-facing, first sheet) ----------
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_view.showGridLines = False

    ws["A1"] = "Plan My Day — Daily Report"
    ws["A1"].font = Font(name=FONT, bold=True, size=15, color=INK)
    ws["A2"] = f"{user['name']}  ·  {user.get('role','')}  ·  {on_date.strftime('%d %b %Y')}"
    ws["A2"].font = Font(name=FONT, size=10, color=MUTED)

    r = 4
    ws.cell(row=r, column=1, value="Monthly scorecard").font = Font(name=FONT, bold=True, size=11, color=INK)
    r += 1
    sc_headers = ["KPI", "Target", "MTD", "Done %", "Gap", "Need/day", "Status"]
    for c, h in enumerate(sc_headers, start=1):
        cell = ws.cell(row=r, column=c, value=h)
        cell.font = Font(name=FONT, bold=True, color="FFFFFF", size=10)
        cell.fill = PatternFill("solid", fgColor=HEAD_FILL)
        cell.border = BORDER
    for i, s in enumerate(scards, start=1):
        vals = [s["kpi_name"], s["monthly_target"], s["achieved_mtd"],
                f'{s["achievement_pct"]}%', s["gap"], s["required_run_rate"], s["status"]]
        band = PatternFill("solid", fgColor=BAND) if i % 2 == 0 else None
        for c, v in enumerate(vals, start=1):
            cell = ws.cell(row=r + i, column=c, value=v)
            cell.font = Font(name=FONT, size=10, color=INK)
            cell.border = BORDER
            if band:
                cell.fill = band
        st_cell = ws.cell(row=r + i, column=len(vals))
        st_cell.font = Font(name=FONT, size=10, bold=True, color=_status_color(s["status"]))
    r = r + len(scards) + 2

    # task + alignment summary
    done = sum(1 for t in tasks if t.get("status") == "Done")
    unlinked = sum(1 for t in tasks if str(t.get("goal_aligned")) == "No")
    builds = sum(1 for t in tasks if str(t.get("horizon")) == "Build")
    ws.cell(row=r, column=1, value="Today at a glance").font = Font(name=FONT, bold=True, size=11, color=INK)
    r += 1
    glance = [
        ("Tasks planned", len(tasks)),
        ("Completed", done),
        ("Build / prep tasks", builds),
        ("Unlinked to a goal", unlinked),
    ]
    for label, val in glance:
        ws.cell(row=r, column=1, value=label).font = Font(name=FONT, size=10, color=MUTED)
        ws.cell(row=r, column=2, value=val).font = Font(name=FONT, size=10, bold=True, color=INK)
        r += 1
    r += 1

    if upd:
        for label, key in [("What worked", "what_worked"),
                           ("What didn't", "what_did_not_work"),
                           ("Numbers update", "numbers_update")]:
            txt = str(upd.get(key, "") or "").strip()
            if txt:
                ws.cell(row=r, column=1, value=label).font = Font(name=FONT, bold=True, size=10, color=INK)
                cell = ws.cell(row=r, column=2, value=txt)
                cell.font = Font(name=FONT, size=10, color=INK)
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=7)
                r += 1

    for col, w in zip("ABCDEFG", [22, 12, 10, 9, 10, 10, 12]):
        ws.column_dimensions[col].width = w

    # ---------- meta (contract header) ----------
    ws_meta = wb.create_sheet("meta")
    meta_row = {
        "report_version": REPORT_VERSION, "app_version": APP_VERSION,
        "user_key": user_key, "name": user["name"], "role": user.get("role", ""),
        "department": user.get("department", ""), "report_date": date_str,
        "generated_at": dt.datetime.now().isoformat(timespec="seconds"),
        "month": month,
    }
    _write_table(ws_meta, META_COLS, [meta_row])
    _autofit(ws_meta)

    # ---------- scorecard (data) ----------
    ws_sc = wb.create_sheet("scorecard")
    _write_table(ws_sc, SCORECARD_COLS, scards)
    _autofit(ws_sc)

    # ---------- tasks (data) ----------
    ws_t = wb.create_sheet("tasks")
    _write_table(ws_t, TASK_COLS, tasks)
    _autofit(ws_t)

    # ---------- day_update (data) ----------
    ws_u = wb.create_sheet("day_update")
    _write_table(ws_u, UPDATE_COLS, [upd] if upd else [])
    _autofit(ws_u)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"PlanMyDay_{user_key}_{date_str}.xlsx"
    return filename, buf.getvalue()


# ============================================================================
# Desktop artifacts: write the daily brief (.md) and the report (.docx) to disk
# inside the user's workspace folders. The app calls these instead of streaming
# a download, since on desktop the files live on the D: drive.
# ============================================================================

def _gather(user_key, on_date):
    """Shared data pull for md/docx so the three formats never disagree."""
    user = storage.get_user(user_key)
    if isinstance(on_date, str):
        on_date = dt.date.fromisoformat(on_date)
    month = on_date.strftime("%Y-%m")
    return {
        "user": user,
        "on_date": on_date,
        "date_str": on_date.isoformat(),
        "month": month,
        "scards": _scorecard_rows(user_key, month, on_date),
        "tasks": _task_rows(user_key, on_date.isoformat()),
        "upd": _day_update(user_key, on_date.isoformat()),
    }


def write_brief_md(user_key, report_date=None):
    """Write the day's brief as Markdown into <user>/briefs/. Returns the path."""
    d = _gather(user_key, report_date or dt.date.today())
    u, sc, tasks, upd = d["user"], d["scards"], d["tasks"], d["upd"]

    worst = nudge.worst_status(sc) if sc else None
    lines = []
    lines.append(f"# Daily Brief — {u['name']}")
    lines.append(f"_{u.get('role','')} · {d['on_date'].strftime('%d %b %Y')}_\n")

    lines.append("## Monthly scorecard")
    lines.append("| KPI | Target | MTD | Done % | Gap | Need/day | Status |")
    lines.append("|---|---|---|---|---|---|---|")
    for s in sc:
        lines.append(f"| {s['kpi_name']} | {_fmt(s['monthly_target'])} | {_fmt(s['achieved_mtd'])} | "
                     f"{s['achievement_pct']}% | {_fmt(s['gap'])} | {_fmt(s['required_run_rate'])} | {s['status']} |")
    lines.append("")

    if worst and worst["status"] in ("Behind", "Critical"):
        lines.append(f"> **Focus:** {worst['kpi_name']} is {worst['status'].lower()} — "
                     f"lead today with work that moves it.\n")

    done = sum(1 for t in tasks if t.get("status") == "Done")
    unlinked = sum(1 for t in tasks if str(t.get("goal_aligned")) == "No")
    builds = sum(1 for t in tasks if str(t.get("horizon")) == "Build")
    lines.append("## Today")
    lines.append(f"- Planned: {len(tasks)} · Completed: {done} · "
                 f"Build/prep: {builds} · Unlinked: {unlinked}")
    if unlinked:
        lines.append(f"- ⚠️ {unlinked} task(s) don't link to a goal — kept, but flagged.")
    lines.append("")

    if tasks:
        lines.append("## Tasks")
        for t in tasks:
            tag = t.get("horizon", "") or ""
            al = "" if str(t.get("goal_aligned")) != "No" else " · _unlinked_"
            lines.append(f"- **[{t.get('priority','')}·{tag}]** {t.get('title','')}"
                         f" → {t.get('linked_kpi','') or '(no KPI)'}{al}")
        lines.append("")

    if upd:
        for label, key in [("What worked", "what_worked"),
                           ("What didn't", "what_did_not_work"),
                           ("Numbers", "numbers_update")]:
            v = str(upd.get(key, "") or "").strip()
            if v:
                lines.append(f"**{label}:** {v}\n")

    os.makedirs(paths.user_briefs_dir(user_key), exist_ok=True)
    path = os.path.join(paths.user_briefs_dir(user_key), f"brief_{d['date_str']}.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return path


def write_report_docx(user_key, report_date=None):
    """Write the day's report as a Word .docx into <user>/reports/. Returns the path."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = _gather(user_key, report_date or dt.date.today())
    u, sc, tasks, upd = d["user"], d["scards"], d["tasks"], d["upd"]

    doc = Document()
    # python-docx ships a <w:zoom> with no percent attr, which fails strict OOXML
    # validation; set it so every generated file is clean.
    from docx.oxml.ns import qn as _qn
    _z = doc.settings.element.find(_qn("w:zoom"))
    if _z is not None and _z.get(_qn("w:percent")) is None:
        _z.set(_qn("w:percent"), "100")
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    title = doc.add_heading("Plan My Day — Daily Report", level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(f"{u['name']}  ·  {u.get('role','')}  ·  {d['on_date'].strftime('%d %b %Y')}")
    run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

    doc.add_heading("Monthly scorecard", level=1)
    t = doc.add_table(rows=1, cols=7)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["KPI", "Target", "MTD", "Done %", "Gap", "Need/day", "Status"]):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
    for s in sc:
        row = t.add_row().cells
        vals = [s["kpi_name"], _fmt(s["monthly_target"]), _fmt(s["achieved_mtd"]),
                f'{s["achievement_pct"]}%', _fmt(s["gap"]), _fmt(s["required_run_rate"]), s["status"]]
        for i, v in enumerate(vals):
            row[i].text = str(v)

    done = sum(1 for x in tasks if x.get("status") == "Done")
    unlinked = sum(1 for x in tasks if str(x.get("goal_aligned")) == "No")
    builds = sum(1 for x in tasks if str(x.get("horizon")) == "Build")
    doc.add_heading("Today at a glance", level=1)
    doc.add_paragraph(f"Planned {len(tasks)}  ·  Completed {done}  ·  "
                      f"Build/prep {builds}  ·  Unlinked {unlinked}")

    if tasks:
        doc.add_heading("Tasks", level=1)
        for x in tasks:
            tag = x.get("horizon", "") or ""
            flag = "" if str(x.get("goal_aligned")) != "No" else "  [unlinked]"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{x.get('priority','')}·{tag}] ").bold = True
            p.add_run(f"{x.get('title','')} → {x.get('linked_kpi','') or '(no KPI)'}{flag}")

    if upd:
        doc.add_heading("Day update", level=1)
        for label, key in [("What worked", "what_worked"),
                           ("What didn't work", "what_did_not_work"),
                           ("Numbers moved", "numbers_update"),
                           ("Remarks", "remarks")]:
            v = str(upd.get(key, "") or "").strip()
            if v:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(v)

    os.makedirs(paths.user_reports_dir(user_key), exist_ok=True)
    path = os.path.join(paths.user_reports_dir(user_key), f"PlanMyDay_{user_key}_{d['date_str']}.docx")
    doc.save(path)
    return path


def write_report_xlsx(user_key, report_date=None):
    """Write the xlsx report (same as build_report) to <user>/reports/. Returns path."""
    fname, data = build_report(user_key, report_date)
    os.makedirs(paths.user_reports_dir(user_key), exist_ok=True)
    path = os.path.join(paths.user_reports_dir(user_key), fname)
    with open(path, "wb") as f:
        f.write(data)
    return path


def write_all(user_key, report_date=None):
    """Generate md brief + docx + xlsx for the day. Returns dict of paths."""
    return {
        "brief_md": write_brief_md(user_key, report_date),
        "report_docx": write_report_docx(user_key, report_date),
        "report_xlsx": write_report_xlsx(user_key, report_date),
    }
