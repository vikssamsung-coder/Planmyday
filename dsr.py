"""Daily Status Report (DSR) builder — Word (.docx) output.

Assembles the whole day — targets vs achievement, tasks and the companion's cues with
their outcomes, meetings, follow-ups/reminders, monthly standing, and proven rules —
into one polished, self-documenting Word document the user downloads. Uses python-docx
(already a dependency); returns the .docx as bytes for st.download_button.
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import storage
import nudge

SLATE = "2D4A5E"
AMBER = "E8833A"
INK = RGBColor(0x1B, 0x27, 0x33)
MUTE = RGBColor(0x5C, 0x6B, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _cell_text(cell, text, *, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _blurb(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = True
    r.font.size = Pt(8.5); r.font.color.rgb = MUTE
    p.paragraph_format.space_after = Pt(4)


def _heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt(12); r.font.color.rgb = INK
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), "E7E3DC")
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)


def _table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], SLATE)
        _cell_text(hdr[i], h, bold=True, color=WHITE, size=9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            _cell_text(cells[i], val, size=9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    return t


def build_docx(user, date, month):
    uk = user["user_key"]
    name = user.get("name", uk)
    role = user.get("role", "").replace("_", " ").title()
    y, m = int(date[:4]), int(date[5:7])

    targets = storage.get_targets(uk, month)
    prog = storage.get_monthly_progress(uk, month=month)
    prog_today = prog[prog["date"] == date] if not prog.empty else prog
    tasks = storage.get_tasks(uk, date)
    meetings = storage.get_meetings(uk)
    meetings = meetings[meetings["date"] == date] if not meetings.empty else meetings
    outbox = storage.get_outbox(uk, date)
    outcomes = storage.get_outcomes(uk, date)
    rules = storage.get_rules(uk)
    day_goals = storage.get_day_goals(uk, date)

    done_n = int((tasks["status"] == "Done").sum()) if not tasks.empty else 0
    total_n = int((~tasks["status"].isin(["Dropped"])).sum()) if not tasks.empty else 0
    behind = 0
    for _, r in targets.iterrows():
        s = nudge.score_kpi(r["monthly_target"], r["achieved_mtd"], y, m)
        if s["status"] in ("Behind", "Critical"):
            behind += 1

    doc = Document()
    normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9); section.right_margin = Inches(0.9)

    k = doc.add_paragraph(); kr = k.add_run("PLAN MY DAY")
    kr.bold = True; kr.font.size = Pt(9); kr.font.color.rgb = RGBColor(0x8A, 0x56, 0x21)
    title = doc.add_paragraph(); tr = title.add_run("Daily Status Report")
    tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = INK
    title.paragraph_format.space_after = Pt(2)
    who = doc.add_paragraph(); wr = who.add_run(f"{name}  \u00b7  {role}")
    wr.font.size = Pt(11); wr.font.color.rgb = INK
    dline = doc.add_paragraph()
    nice_date = datetime.strptime(date, "%Y-%m-%d").strftime("%A, %d %B %Y")
    dr = dline.add_run(f"{nice_date}   \u00b7   generated {datetime.now().strftime('%d %b %Y, %H:%M')}")
    dr.font.size = Pt(9); dr.font.color.rgb = MUTE
    rule_p = doc.add_paragraph(); pPr = rule_p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), AMBER)
    pbdr.append(bottom); pPr.append(pbdr)

    _table(doc, ["Tasks done", "Meetings", "Reminders prepared", "KPIs behind"],
           [[f"{done_n}/{total_n}", str(len(meetings)), str(len(outbox)), str(behind)]],
           widths=[1.7, 1.7, 1.7, 1.7])

    _heading(doc, "Today's targets vs achievement")
    _blurb(doc, "What you aimed for today (from your daily goal sheet) vs what you achieved.")
    rows = []
    for _, r in targets.iterrows():
        kpi = r["kpi_name"]; tgt = ""
        for g in day_goals:
            if g["heading"] and nudge.goal_served(g["heading"], [kpi]):
                tgt = g["target_number"]; break
        ach = ""
        if not prog_today.empty:
            hit = prog_today[prog_today["kpi_name"] == kpi]
            if not hit.empty:
                ach = hit.iloc[0]["achieved"]
        rows.append([kpi, tgt or "\u2014", ach or "\u2014"])
    _table(doc, ["KPI", "Target", "Achievement"], rows or [["No targets set", "", ""]],
           widths=[3.4, 1.5, 1.5])

    _heading(doc, "Tasks & coaching")
    _blurb(doc, "Each task, the companion's cue (how to do it well), and \u2014 if you closed the "
                "day \u2014 whether you tried it and how it went.")
    oc_by_task = {r["task_id"]: r for _, r in outcomes.iterrows()} if not outcomes.empty else {}
    trows = []
    if not tasks.empty:
        for _, t in tasks.iterrows():
            if t["status"] == "Dropped":
                continue
            steps = storage.get_task_steps(uk, t["task_id"])
            sd = f"{sum(1 for s in steps if s['done'])}/{len(steps)}" if steps else "\u2014"
            oc = oc_by_task.get(t["task_id"])
            outcome = f"{oc['tried']}/{oc['result']}" if oc is not None and oc["tried"] else "\u2014"
            trows.append([t["title"], t["day_goal"] or "\u2014", t["status"], sd,
                          t["coach_cue"] or "", outcome])
    _table(doc, ["Task", "Goal", "Status", "Steps", "Cue", "Tried/Result"],
           trows or [["No tasks", "", "", "", "", ""]],
           widths=[1.7, 1.1, 0.8, 0.6, 2.1, 0.9])

    _heading(doc, "Meetings logged")
    _blurb(doc, "Conversations recorded today, with the agreed next action.")
    tl = {"new_partner": "New partner", "existing_partner": "Existing partner",
          "client": "Client", "internal": "Internal"}
    mrows = []
    for _, mtg in meetings.iterrows():
        nxt = f"{mtg['next_action']} (by {mtg['next_date']})" if mtg["next_date"] else "\u2014"
        mrows.append([tl.get(mtg["meeting_type"], mtg["meeting_type"]),
                      mtg["identity_value"], mtg["outcome"] or "\u2014", nxt])
    _table(doc, ["Type", "Who", "Outcome", "Next action"],
           mrows or [["No meetings logged", "", "", ""]],
           widths=[1.4, 1.3, 1.8, 1.9])

    _heading(doc, "Reminders prepared")
    _blurb(doc, "Messages your recurring reminders queued for today (sending is manual).")
    orows = [[o["recipient_name"], o["recipient_mobile"], o["status"]]
             for _, o in outbox.iterrows()] if not outbox.empty else []
    _table(doc, ["Recipient", "Mobile", "Status"],
           orows or [["None due today", "", ""]], widths=[2.6, 2.0, 1.4])

    _heading(doc, "Monthly standing (MIS)")
    _blurb(doc, "Where the month stands per KPI \u2014 the North Star every task ladders up to.")
    srows = []
    for _, r in targets.iterrows():
        s = nudge.score_kpi(r["monthly_target"], r["achieved_mtd"], y, m)
        srows.append([r["kpi_name"], str(r["monthly_target"]), str(r["achieved_mtd"]),
                      str(s["gap"]), f"{s['required_run_rate']:.0f}/day", s["status"]])
    _table(doc, ["KPI", "Target", "Achieved", "Gap", "Run-rate", "Status"],
           srows or [["No targets", "", "", "", "", ""]],
           widths=[2.0, 0.9, 0.9, 0.8, 1.0, 1.0])

    _heading(doc, "What's working for you (proven rules)")
    _blurb(doc, "Tactics that have worked before, saved from your end-of-day check-ins. "
                "The companion leads with these next time.")
    rrows = [[r["topic"], r["rule_text"], str(r["successes"]), r["status"]]
             for _, r in rules.iterrows()] if not rules.empty else []
    _table(doc, ["Topic", "Rule", "Successes", "Status"],
           rrows or [["None yet \u2014 close a few days to build these", "", "", ""]],
           widths=[1.4, 3.0, 0.9, 1.0])

    f = doc.add_paragraph(); f.paragraph_format.space_before = Pt(16)
    fr = f.add_run("Generated by Plan My Day \u00b7 companion coach. Targets are from the daily "
                   "goal sheet; achievement entered at day close. Cues and outcomes come from "
                   "the end-of-day check-in. Reflects local data at generation time.")
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0x8A, 0x93, 0xA0)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def docx_to_text(docx_bytes):
    """Plain-text rendering of a DSR .docx — for saving the report's content to the
    cloud/sheet (a Word file can't live in a sheet cell, but its text can)."""
    import io
    from docx import Document
    d = Document(io.BytesIO(docx_bytes))
    lines = []
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)
